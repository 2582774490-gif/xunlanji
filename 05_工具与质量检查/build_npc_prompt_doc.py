from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "雾隐原NPC绘图提示词.md"
OUTPUT = ROOT / "雾隐原首批NPC_GPT可直接输入绘图提示词.docx"
MARKDOWN_OUTPUT = ROOT / "雾隐原首批NPC_GPT可直接输入绘图提示词.md"


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_prompt_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_margins(cell)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    tc_pr.append(shd)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    run = paragraph.add_run(text)
    set_run_font(run, size=9.2, name="Microsoft YaHei")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def extract_code_block(markdown, heading):
    match = re.search(rf"## {re.escape(heading)}\s*\n\n```text\n(.*?)\n```", markdown, re.S)
    if not match:
        raise RuntimeError(f"找不到区块：{heading}")
    return match.group(1).strip()


def extract_characters(markdown):
    rows = []
    in_table = False
    for line in markdown.splitlines():
        if line.startswith("| ID/角色"):
            in_table = True
            continue
        if in_table and line.startswith("|---"):
            continue
        if in_table and line.startswith("|"):
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(parts) == 3:
                key_name = parts[0].replace("`", "")
                match = re.match(r"(P-N\d+)\s+(.+)", key_name)
                if match:
                    rows.append({
                        "id": match.group(1),
                        "name": match.group(2),
                        "turnaround": parts[1],
                        "face": parts[2],
                    })
        elif in_table:
            break
    if len(rows) != 15:
        raise RuntimeError(f"应读取 15 位角色，实际读取 {len(rows)} 位")
    return rows


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, (46, 116, 181), 18, 10),
        ("Heading 2", 13, (46, 116, 181), 14, 7),
        ("Heading 3", 12, (31, 77, 120), 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = header_p.add_run("《寻岚记》｜雾隐原 NPC 绘图提示词")
    set_run_font(hrun, size=8.5, color=(90, 90, 90))
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer_p.add_run("内部制作参考 · 原创角色资产")
    set_run_font(frun, size=8, color=(120, 120, 120))


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("《寻岚记》雾隐原首批 NPC")
    set_run_font(run, size=25, bold=True, color=(31, 77, 120))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(26)
    run2 = p2.add_run("可直接输入 GPT 的四视图与面部细节绘图提示词")
    set_run_font(run2, size=13, color=(82, 82, 82))
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(18)
    run3 = p3.add_run("v0.2｜15 位命名 NPC｜先四视图，后面部锚点图")
    set_run_font(run3, size=9.5, color=(105, 105, 105))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(20)
    note.paragraph_format.space_after = Pt(0)
    note.add_run("使用原则：").bold = True
    note.add_run("采用原创中国 3D 修仙动画建模感，不复制任何现有作品的角色、服装、纹样或镜头。")
    set_paragraph_shading(note, "F4F6F9")
    for run in note.runs:
        set_run_font(run, size=10)


def add_usage(doc):
    doc.add_heading("怎么使用", level=1)
    steps = [
        "每次只复制一位 NPC 的“完整四视图提示词”，粘贴到 GPT 图像生成对话中发送。",
        "从生成结果中选出四张脸、服装和比例最一致的一张，保存为该角色的“四视图锚点图”。",
        "上传锚点图到 GPT，再复制同一角色的“完整面部细节提示词”发送。",
        "若四视图不一致，不要直接生成面部图；先重新生成四视图，并在资产卡记录最终采用的版本。",
    ]
    for index, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(step)
        set_run_font(run, size=10.5)

    warning = doc.add_paragraph()
    warning.paragraph_format.space_before = Pt(8)
    warning.paragraph_format.space_after = Pt(4)
    run = warning.add_run("注意：")
    set_run_font(run, size=10.5, bold=True, color=(122, 90, 0))
    run2 = warning.add_run("面部细节提示词必须在已上传对应四视图锚点图之后使用。")
    set_run_font(run2, size=10.5, color=(122, 90, 0))


def add_prompt_sections(doc, characters, universal_turnaround, universal_face):
    doc.add_page_break()
    doc.add_heading("一、完整四视图提示词（直接复制）", level=1)
    intro = doc.add_paragraph("以下每段已经包含全部通用规范和角色设定。一次只发送一段。")
    for run in intro.runs:
        set_run_font(run, size=10.5, color=(82, 82, 82))

    for character in characters:
        doc.add_heading(f"{character['id']}｜{character['name']}", level=2)
        add_prompt_block(doc, f"{universal_turnaround}\n\n{character['turnaround']}")

    doc.add_page_break()
    doc.add_heading("二、完整面部细节提示词（上传锚点图后复制）", level=1)
    intro2 = doc.add_paragraph("每段使用前，请先上传同一角色已确认的四视图锚点图，并让它成为图像 1。")
    for run in intro2.runs:
        set_run_font(run, size=10.5, color=(82, 82, 82))

    for character in characters:
        doc.add_heading(f"{character['id']}｜{character['name']}", level=2)
        add_prompt_block(doc, f"{universal_face}\n\n{character['face']}")


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    universal_turnaround = extract_code_block(markdown, "通用四视图提示词")
    universal_face = extract_code_block(markdown, "通用面部细节提示词")
    characters = extract_characters(markdown)

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_usage(doc)
    add_prompt_sections(doc, characters, universal_turnaround, universal_face)
    doc.core_properties.title = "《寻岚记》雾隐原首批 NPC 绘图提示词"
    doc.core_properties.subject = "可直接输入 GPT 的角色四视图与面部细节提示词"
    doc.core_properties.author = "《寻岚记》项目"
    doc.save(OUTPUT)
    lines = [
        "# 《寻岚记》雾隐原首批 NPC：可直接输入 GPT 的绘图提示词",
        "",
        "先复制一位 NPC 的完整四视图提示词并发送。选定一致的四视图后，将其上传为图像 1，再复制同一角色的面部细节提示词发送。所有提示词均采用原创中国 3D 修仙动画建模感，不复刻任何现有作品。",
        "",
        "## 一、完整四视图提示词",
        "",
    ]
    for character in characters:
        lines += [
            f"### {character['id']}｜{character['name']}",
            "",
            "```text",
            universal_turnaround,
            "",
            character["turnaround"],
            "```",
            "",
        ]
    lines += ["## 二、完整面部细节提示词（请先上传对应四视图锚点图）", ""]
    for character in characters:
        lines += [
            f"### {character['id']}｜{character['name']}",
            "",
            "```text",
            universal_face,
            "",
            character["face"],
            "```",
            "",
        ]
    MARKDOWN_OUTPUT.write_text("\n".join(lines), encoding="utf-8")
    print(OUTPUT)
    print(MARKDOWN_OUTPUT)


if __name__ == "__main__":
    main()
